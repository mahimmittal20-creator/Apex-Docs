from docx import Document
from docx.shared import Inches, Pt, Twips
from typing import Optional
import os
import pythoncom
from .models import Resume, Experience, Education
from .utils import process_bullet_points

def convert_word_to_pdf_hidden(word_path: str, pdf_path: str):
    """
    Convert Word document to PDF using Word COM with hidden window.
    This prevents Word from flashing on screen during conversion.
    """
    pythoncom.CoInitialize()
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  # Hide Word window
        word.DisplayAlerts = False  # Suppress alerts
        
        # Convert to absolute path
        abs_word_path = os.path.abspath(word_path)
        abs_pdf_path = os.path.abspath(pdf_path)
        
        doc = word.Documents.Open(abs_word_path)
        doc.SaveAs(abs_pdf_path, FileFormat=17)  # 17 = PDF format
        doc.Close()
        word.Quit()
    finally:
        pythoncom.CoUninitialize()

# =============================================================================
# FONT CONFIGURATION - Change these values to customize your resume
# =============================================================================

# Font family - Common options:
# "Calibri", "Arial", "Times New Roman", "Garamond", "Georgia", 
# "Cambria", "Helvetica", "Verdana", "Tahoma", "Century Gothic"
FONT_NAME = "Calibri"

# Font sizes (in points)
FONT_SIZE_BULLET = 11  # Font size for experience bullet points
FONT_SIZE_SKILLS = 10  # Font size for skills
FONT_SIZE_SUMMARY = 10  # Font size for summary
FONT_SIZE_CERTIFICATIONS = 10  # Font size for certifications
FONT_SIZE_JOB_TITLE = 12  # Font size for job titles
FONT_SIZE_DATE = 11  # Font size for dates
FONT_SIZE_SECTION_HEADER = 11  # Font size for section headers (Summary, Skills, etc.)
FONT_SIZE_NAME = 19  # Font size for name
FONT_SIZE_CONTACT = 9  # Font size for contact info

# Bullet indentation (in inches)
BULLET_INDENT = 0.25  # Left indent for bullet text
BULLET_HANG = 0.11    # Hanging indent (bullet hangs to the left of text)

def apply_font(run, size=None, bold=False):
    """Apply font name and optionally size and bold to a run."""
    run.font.name = FONT_NAME
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    return run

def apply_bullet_format(paragraph, text, font_size):
    """
    Apply bullet formatting with proper hanging indent.
    This ensures wrapped text aligns with the first letter, not under the bullet.
    """
    paragraph.clear()
    # Set indentation: left indent for the text, negative first line for hanging bullet
    paragraph.paragraph_format.left_indent = Inches(BULLET_INDENT)
    paragraph.paragraph_format.first_line_indent = Inches(-BULLET_HANG)
    # Add bullet and text
    run = paragraph.add_run(f"• {text}")
    apply_font(run, font_size)
    return paragraph

def add_bullet_paragraph(document, text, font_size=FONT_SIZE_BULLET):
    """
    Add a new paragraph with proper bullet formatting.
    """
    paragraph = document.add_paragraph()
    apply_bullet_format(paragraph, text, font_size)
    return paragraph

def format_skills_with_bold_category(paragraph, skills, font_size=None):
    """
    Format skills with bold category names.
    Expected format: "Category Name: skill1, skill2, skill3"
    """
    if font_size is None:
        font_size = FONT_SIZE_SKILLS
    paragraph.clear()
    for i, skill_line in enumerate(skills):
        if ':' in skill_line:
            # Split into category and skills
            category, skill_list = skill_line.split(':', 1)
            # Add bold category
            bold_run = paragraph.add_run(category.strip() + ': ')
            apply_font(bold_run, font_size, bold=True)
            # Add regular skills
            regular_run = paragraph.add_run(skill_list.strip())
            apply_font(regular_run, font_size)
        else:
            # No category, just add as is
            run = paragraph.add_run(skill_line)
            apply_font(run, font_size)
        
        # Add newline if not last skill
        if i < len(skills) - 1:
            paragraph.add_run('\n')

def generate_pdf_resume(resume: Resume) -> str:
    """
    Generate PDF by first creating Word document, then converting to PDF.
    This ensures PDF looks exactly like the Word file.
    """
    # First generate the Word document
    word_path = generate_word_resume(resume)
    
    # Convert Word to PDF (hidden, no flashing window)
    pdf_path = word_path.replace('.docx', '.pdf')
    convert_word_to_pdf_hidden(word_path, pdf_path)
    
    return pdf_path

def generate_word_resume(resume: Resume) -> str:
    template_path = "backend/templates/resume_template.docx"
    
    # Ensure all experience entries are present for robust templating
    webkorps_exp = next((exp for exp in resume.experience if exp.company == "WebKorps"), None)
    ibm_exp = next((exp for exp in resume.experience if exp.company == "IBM"), None)
    americankorps_exp = next((exp for exp in resume.experience if exp.company == "AmericanKorps"), None)

    if not os.path.exists(template_path):
        # Fallback to basic generation if template is missing or not a file
        document = Document()
        document.add_heading(resume.name, level=1)
        
        contact_info = f"{resume.email}"
        if resume.phone: contact_info += f" | {resume.phone}"
        if resume.location: contact_info += f" | {resume.location}"
        document.add_paragraph(contact_info)

        links_info = []
        if resume.linkedin: links_info.append(f"LinkedIn: {resume.linkedin}")
        if resume.github: links_info.append(f"GitHub: {resume.github}")
        if links_info:
            document.add_paragraph(" | ".join(links_info))

        document.add_heading('Summary', level=2)
        summary_para = document.add_paragraph()
        run = summary_para.add_run(resume.summary)
        apply_font(run, FONT_SIZE_SUMMARY)

        if resume.skills:
            document.add_heading('Skills', level=2)
            skills_para = document.add_paragraph()
            format_skills_with_bold_category(skills_para, resume.skills, FONT_SIZE_SKILLS)

        if resume.experience:
            document.add_heading('Experience', level=2)
            for exp in resume.experience:
                document.add_paragraph(f"{exp.title} at {exp.company} ({exp.start_date} - {exp.end_date or 'Present'})")
                description_points = process_bullet_points(exp.description)
                for point in description_points:
                    add_bullet_paragraph(document, point, FONT_SIZE_BULLET)

        if resume.education:
            document.add_heading('Education', level=2)
            for edu in resume.education:
                document.add_paragraph(f"{edu.degree} in {edu.major}, {edu.university} ({edu.graduation_date})")

        if resume.projects:
            document.add_heading('Projects', level=2)
            for project in resume.projects:
                add_bullet_paragraph(document, project, FONT_SIZE_BULLET)
        
        if resume.certifications:
            document.add_heading('Certifications', level=2)
            for cert in resume.certifications:
                add_bullet_paragraph(document, cert, FONT_SIZE_CERTIFICATIONS)
    else:
        document = Document(template_path)

        # Collect all paragraphs first to avoid iteration issues
        paragraphs_list = list(document.paragraphs)
        description_paragraphs_to_replace = []
        
        # Helper function to check if placeholder exists in paragraph (handles split runs)
        def has_placeholder(para, placeholder):
            # Check full paragraph text
            if placeholder in para.text:
                return True
            # Check if placeholder is split across runs
            full_text = "".join([run.text for run in para.runs])
            if placeholder in full_text:
                return True
            return False

        for paragraph in paragraphs_list:
            # Replace top-level contact info, summary, skills
            if "{{name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{name}}", resume.name)
            if "{{email}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{email}}", resume.email)
            if "{{phone}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{phone}}", resume.phone if resume.phone else "")
            if "{{location}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{location}}", resume.location if resume.location else "")
            if "{{summary}}" in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run(resume.summary)
                apply_font(run, FONT_SIZE_SUMMARY)
            if "{{skills}}" in paragraph.text:
                # Use the new formatting function for skills with bold categories
                format_skills_with_bold_category(paragraph, resume.skills, FONT_SIZE_SKILLS)
            
            # --- Handle Experience Section ---
            # Replace titles with font size and include dates
            if "{{webkorps_title}}" in paragraph.text:
                paragraph.clear()
                if webkorps_exp:
                    # Add title (bold)
                    title_run = paragraph.add_run(webkorps_exp.title)
                    apply_font(title_run, FONT_SIZE_JOB_TITLE, bold=True)
                    # Add tab and date
                    date_text = f"\t\t{webkorps_exp.start_date}-{webkorps_exp.end_date or 'Present'}"
                    date_run = paragraph.add_run(date_text)
                    apply_font(date_run, FONT_SIZE_DATE)
            
            if "{{ibm_title}}" in paragraph.text:
                paragraph.clear()
                if ibm_exp:
                    # Add title (bold)
                    title_run = paragraph.add_run(ibm_exp.title)
                    apply_font(title_run, FONT_SIZE_JOB_TITLE, bold=True)
                    # Add tab and date
                    date_text = f"\t\t{ibm_exp.start_date}-{ibm_exp.end_date or 'Present'}"
                    date_run = paragraph.add_run(date_text)
                    apply_font(date_run, FONT_SIZE_DATE)
            
            if "{{americankorps_title}}" in paragraph.text:
                paragraph.clear()
                if americankorps_exp:
                    # Add title (bold)
                    title_run = paragraph.add_run(americankorps_exp.title)
                    apply_font(title_run, FONT_SIZE_JOB_TITLE, bold=True)
                    # Add tab and date
                    date_text = f"\t\t{americankorps_exp.start_date}-{americankorps_exp.end_date or 'Present'}"
                    date_run = paragraph.add_run(date_text)
                    apply_font(date_run, FONT_SIZE_DATE)
            
            # WebKorps Description - mark for replacement
            if has_placeholder(paragraph, "{{webkorps_description}}") or "{{webkorps_description}}" in paragraph.text:
                description_paragraphs_to_replace.append(("webkorps", paragraph, webkorps_exp))
            
            # IBM Description
            if has_placeholder(paragraph, "{{ibm_description}}") or "{{ibm_description}}" in paragraph.text:
                description_paragraphs_to_replace.append(("ibm", paragraph, ibm_exp))
            
            # AmericanKorps Description
            if has_placeholder(paragraph, "{{americankorps_description}}") or "{{americankorps_description}}" in paragraph.text:
                description_paragraphs_to_replace.append(("americankorps", paragraph, americankorps_exp))
            
            # --- Handle Certifications Section ---
            if "{{certifications}}" in paragraph.text or "• {{certifications}}" in paragraph.text:
                # Get parent for inserting additional certification paragraphs
                parent = paragraph._element.getparent()
                para_index = parent.index(paragraph._element)
                
                if resume.certifications:
                    # First certification uses the existing paragraph
                    apply_bullet_format(paragraph, resume.certifications[0], FONT_SIZE_CERTIFICATIONS)
                    
                    # Add remaining certifications as new paragraphs
                    for i, cert in enumerate(resume.certifications[1:], 1):
                        cert_para = add_bullet_paragraph(document, cert, FONT_SIZE_CERTIFICATIONS)
                        # Move to correct position
                        cert_para._element.getparent().remove(cert_para._element)
                        parent.insert(para_index + i, cert_para._element)
                else:
                    paragraph.clear()

        # Now handle description paragraphs - create separate paragraphs for each bullet
        for exp_type, para, exp in description_paragraphs_to_replace:
            if exp:
                description_points = process_bullet_points(exp.description)
                if description_points:
                    # Get parent element and index for inserting new paragraphs
                    parent = para._element.getparent()
                    para_index = parent.index(para._element)
                    
                    # Set first paragraph to bullet style
                    apply_bullet_format(para, description_points[0], FONT_SIZE_BULLET)
                    
                    # Create and insert remaining bullets as separate paragraphs
                    # We need to insert in reverse order to maintain correct positions
                    for i in range(len(description_points) - 1, 0, -1):
                        point = description_points[i]
                        # Create new paragraph with bullet style
                        temp_para = add_bullet_paragraph(document, point, FONT_SIZE_BULLET)
                        # Remove from end and insert at correct position
                        temp_para._element.getparent().remove(temp_para._element)
                        parent.insert(para_index + 1, temp_para._element)
            else:
                para.clear()


    file_path = f"./tmp/{resume.name.replace(' ', '_')}_resume.docx"
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    document.save(file_path)
    return file_path